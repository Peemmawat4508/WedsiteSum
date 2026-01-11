from fastapi import FastAPI, UploadFile, File, Depends, HTTPException, status, Request
from fastapi.responses import Response
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session
from datetime import datetime, timedelta
from typing import Optional
import os
import PyPDF2
import pdfplumber
from pydantic import BaseModel
import uvicorn
from dotenv import load_dotenv
import json

# Load environment variables from .env file
load_dotenv()

from database import SessionLocal, engine, Base
from models import Document
from schemas import DocumentResponse, SummaryResponse, QueryRequest, QueryResponse, ChatRequest, ChatResponse, ImageGenerationRequest, ImageGenerationResponse, ExportRequest, GrammarCheckRequest, GrammarCheckResponse
from summarizer import summarize_text, create_chunks, create_embeddings, query_documents, generate_rag_answer, chat_with_gpt, generate_image, grammar_check

app = FastAPI(title="Document Summarizer API (Guest Mode)")

# Setup startup event for table creation
@app.on_event("startup")
async def startup_event():
    try:
        # Create tables
        Base.metadata.create_all(bind=engine)
        print("Tables created successfully")
    except Exception as e:
        print(f"Error in startup_event: {e}")

# Global Exception Handler for detailed error reporting
from fastapi.responses import JSONResponse
import traceback

@app.exception_handler(Exception)
async def global_exception_handler(request, exc):
    error_msg = f"Global Error: {str(exc)}\n{traceback.format_exc()}"
    print(error_msg)
    return JSONResponse(
        status_code=500,
        content={"detail": str(exc), "traceback": str(traceback.format_exc())}
    )

@app.get("/debug")
async def debug_endpoint():
    """Diagnostic endpoint to check environment"""
    debug_info = {
        "status": "online",
        "timestamp": datetime.utcnow().isoformat(),
        "env_vars": {
            "VERCEL": os.environ.get("VERCEL"),
            "VITE_GOOGLE_CLIENT_ID": "Set" if os.environ.get("VITE_GOOGLE_CLIENT_ID") else "Not Set",
            "OPENAI_API_KEY": "Set" if os.environ.get("OPENAI_API_KEY") else "Not Set",
        },
        "filesystem": {
            "tmp_exists": os.path.exists("/tmp"),
            "tmp_writable": os.access("/tmp", os.W_OK),
            "db_path": str(engine.url),
            "db_file_exists": os.path.exists("/tmp/documents.db") if "sqlite" in str(engine.url) else "N/A"
        }
    }
    
    # Check DB connection
    try:
        db = SessionLocal()
        doc_count = db.query(Document).count()
        debug_info["database"] = {"status": "connected", "document_count": doc_count}
        db.close()
    except Exception as e:
        debug_info["database"] = {"status": "error", "error": str(e)}
        
    return debug_info

@app.get("/")
async def root():
    return {"status": "ok", "message": "Backend is running in Guest Mode"}

# CORS configuration
cors_origins = os.getenv("CORS_ORIGINS", "http://localhost:3000,http://localhost:5173").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=[origin.strip() for origin in cors_origins],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Middleware to check content size
@app.middleware("http")
async def validate_content_length(request: Request, call_next):
    content_length = request.headers.get("content-length")
    if content_length:
        content_length = int(content_length)
        if content_length > 4.5 * 1024 * 1024:  # 4.5MB Vercel limit
            return JSONResponse(
                status_code=413,
                content={"detail": "File too large. Maximum size is 4.5MB."}
            )
    return await call_next(request)

# Dependency to get DB session
def get_db():
    try:
        db = SessionLocal()
        try:
            yield db
        finally:
            db.close()
    except Exception as e:
        print(f"Error in get_db: {e}")
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Database service is currently unavailable"
        )

# --- Text Extraction Functions ---
# Extract text from PDF - improved for Thai characters
def extract_text_from_pdf(file_content: bytes) -> str:
    import io
    text = ""
    
    # Try pdfplumber first (better for Thai/unicode)
    try:
        pdf_file = io.BytesIO(file_content)
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return text
    except Exception as e:
        print(f"pdfplumber extraction failed: {e}, trying PyPDF2...")
    
    # Fallback to PyPDF2
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        print(f"PyPDF2 extraction failed: {e}")
    
    return text

def extract_text_from_txt(file_content: bytes) -> str:
    return file_content.decode('utf-8')

def extract_text_from_docx(file_content: bytes) -> str:
    import io
    try:
        from docx import Document as DocxDocument
        docx_file = io.BytesIO(file_content)
        doc = DocxDocument(docx_file)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)
    except ImportError:
        raise Exception("python-docx library is required for DOCX files")
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")

def extract_text_from_xlsx(file_content: bytes) -> str:
    import io
    try:
        import openpyxl
        xlsx_file = io.BytesIO(file_content)
        workbook = openpyxl.load_workbook(xlsx_file, data_only=True)
        text = []
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text.append(f"Sheet: {sheet_name}")
            for row in sheet.iter_rows(values_only=True):
                row_text = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                if row_text.strip():
                    text.append(row_text)
            text.append('')  # Empty line between sheets
        return '\n'.join(text)
    except ImportError:
        raise Exception("openpyxl library is required for XLSX files")
    except Exception as e:
        raise Exception(f"Error extracting text from XLSX: {str(e)}")

def extract_text_from_csv(file_content: bytes) -> str:
    import io
    import csv
    try:
        csv_file = io.BytesIO(file_content)
        # Try different encodings
        for encoding in ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']:
            try:
                csv_file.seek(0)
                text = csv_file.read().decode(encoding)
                csv_file.seek(0)
                reader = csv.reader(io.StringIO(text))
                rows = []
                for row in reader:
                    rows.append(' | '.join(row))
                return '\n'.join(rows)
            except (UnicodeDecodeError, Exception):
                continue
        raise Exception("Could not decode CSV file")
    except Exception as e:
        raise Exception(f"Error extracting text from CSV: {str(e)}")

def extract_text_from_md(file_content: bytes) -> str:
    try:
        import markdown
        md_text = file_content.decode('utf-8')
        # Convert markdown to plain text by removing markdown syntax
        html = markdown.markdown(md_text)
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, 'html.parser')
        return soup.get_text()
    except ImportError:
        # Fallback: just decode and return
        return file_content.decode('utf-8')
    except Exception as e:
        return file_content.decode('utf-8', errors='ignore')

def extract_text_from_html(file_content: bytes) -> str:
    try:
        from bs4 import BeautifulSoup
        html_content = file_content.decode('utf-8', errors='ignore')
        soup = BeautifulSoup(html_content, 'html.parser')
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        return soup.get_text(separator='\n', strip=True)
    except ImportError:
        # Fallback: basic text extraction
        return file_content.decode('utf-8', errors='ignore')
    except Exception as e:
        return file_content.decode('utf-8', errors='ignore')

def extract_text_from_image(file_content: bytes, filename: str) -> str:
    try:
        from PIL import Image
        import pytesseract
        import io
        
        # Open image
        image = Image.open(io.BytesIO(file_content))
        
        # Perform OCR
        try:
            text = pytesseract.image_to_string(image, lang='eng+tha')
        except:
            text = pytesseract.image_to_string(image, lang='eng')
        
        return text.strip()
    except ImportError:
        raise Exception("Pillow and pytesseract libraries are required for image OCR.")
    except Exception as e:
        raise Exception(f"Error extracting text from image: {str(e)}")

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Extract text from various file types based on extension."""
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        return extract_text_from_pdf(file_content)
    elif filename_lower.endswith('.txt'):
        return extract_text_from_txt(file_content)
    elif filename_lower.endswith(('.docx', '.doc')):
        return extract_text_from_docx(file_content)
    elif filename_lower.endswith(('.xlsx', '.xls')):
        return extract_text_from_xlsx(file_content)
    elif filename_lower.endswith('.csv'):
        return extract_text_from_csv(file_content)
    elif filename_lower.endswith(('.md', '.markdown')):
        return extract_text_from_md(file_content)
    elif filename_lower.endswith(('.html', '.htm')):
        return extract_text_from_html(file_content)
    elif filename_lower.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp')):
        return extract_text_from_image(file_content, filename)
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type: {filename_lower.split('.')[-1] if '.' in filename_lower else 'unknown'}"
        )

# --- End Text Extraction Functions ---

@app.post("/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    # Read file content
    file_content = await file.read()
    
    # Extract text from file (supports multiple file types)
    try:
        text = extract_text_from_file(file_content, file.filename)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Error processing file: {str(e)}"
        )
    
    if not text.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Could not extract text from document"
        )
    
    # Create chunks and embeddings for RAG
    chunks = create_chunks(text, chunk_size=1000, overlap=200)
    embeddings = []
    
    if chunks:
        embeddings = create_embeddings(chunks)
    
    # Prepare chunks data for storage
    chunks_data = [
        {"text": chunk, "embedding": emb} 
        for chunk, emb in zip(chunks, embeddings) if emb
    ]
    
    # Create document record
    db_document = Document(
        filename=file.filename,
        original_text=text,  # Store full text
        chunks=json.dumps(chunks_data) if chunks_data else None,
        embeddings=json.dumps(embeddings) if embeddings else None
    )
    db.add(db_document)
    db.commit()
    db.refresh(db_document)
    
    return DocumentResponse(
        id=db_document.id,
        filename=db_document.filename,
        uploaded_at=db_document.uploaded_at
    )

@app.post("/summarize/{document_id}", response_model=SummaryResponse)
async def summarize_document(
    document_id: int,
    db: Session = Depends(get_db)
):
    # Get document
    document = db.query(Document).filter(Document.id == document_id).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Get full text from document
    text = document.original_text
    
    if not text:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document has no text content"
        )
    
    # Generate summary using OpenAI with better prompts
    summary = summarize_text(text, max_length=300)
    
    # Update document with summary
    document.summary = summary
    db.commit()
    
    return SummaryResponse(
        document_id=document.id,
        summary=summary,
        filename=document.filename
    )

@app.get("/documents", response_model=list[DocumentResponse])
async def get_documents(db: Session = Depends(get_db)):
    try:
        # Ensure tables exist (helper for Vercel cold starts)
        # In production this is usually redundant but crucial for sqlite in /tmp
        # if the startup event missed it.
        from sqlalchemy import inspect
        inspector = inspect(engine)
        if not inspector.has_table("documents"):
             print("Table 'documents' not found, creating tables...")
             Base.metadata.create_all(bind=engine)

        documents = db.query(Document).order_by(Document.uploaded_at.desc()).all()
        
        return [
            DocumentResponse(
                id=doc.id,
                filename=doc.filename,
                uploaded_at=doc.uploaded_at,
                summary=doc.summary
            )
            for doc in documents
        ]
    except Exception as e:
        print(f"Error getting documents: {e}")
        # Return empty list instead of 500 error if DB is wonky
        return []

@app.delete("/documents/{document_id}")
async def delete_document(
    document_id: int,
    db: Session = Depends(get_db)
):
    """Delete a document"""
    document = db.query(Document).filter(Document.id == document_id).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    db.delete(document)
    db.commit()
    
    return {"message": "Document deleted successfully"}

@app.post("/query", response_model=QueryResponse)
async def query_document(
    query_data: QueryRequest,
    db: Session = Depends(get_db)
):
    """
    Query documents using RAG. 
    """
    # Get documents to search
    if query_data.document_id:
        documents = db.query(Document).filter(Document.id == query_data.document_id).all()
    else:
        documents = db.query(Document).all()
    
    if not documents:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No documents found"
        )
    
    # Collect all chunks from documents
    all_chunks = []
    for doc in documents:
        if doc.chunks:
            try:
                chunks_data = json.loads(doc.chunks)
                # Add document info to each chunk
                for chunk_data in chunks_data:
                    chunk_data['document_id'] = doc.id
                    chunk_data['filename'] = doc.filename
                all_chunks.extend(chunks_data)
            except json.JSONDecodeError:
                continue
    
    if not all_chunks:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No document chunks available for RAG queries."
        )
    
    # Query documents using RAG
    relevant_chunks = query_documents(query_data.query, all_chunks, top_k=5)
    
    if not relevant_chunks:
        return QueryResponse(
            answer="No relevant information found in the documents.",
            document_id=query_data.document_id or documents[0].id,
            filename=documents[0].filename if documents else "Unknown"
        )
    
    # Get the primary document (first relevant chunk's document)
    primary_doc_id = relevant_chunks[0].get('document_id', documents[0].id)
    primary_doc = next((d for d in documents if d.id == primary_doc_id), documents[0])
    
    # Generate answer using RAG
    answer = generate_rag_answer(
        query_data.query, 
        relevant_chunks, 
        primary_doc.filename
    )
    
    return QueryResponse(
        answer=answer,
        document_id=primary_doc.id,
        filename=primary_doc.filename,
        relevant_chunks=[chunk['text'] for chunk in relevant_chunks[:3]]  # Return first 3 chunks
    )

@app.post("/chat", response_model=ChatResponse)
async def chat(chat_data: ChatRequest):
    """
    Chat with GPT like ChatGPT - normal conversation without document context.
    """
    if not chat_data.message.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Message cannot be empty"
        )
    
    # Convert conversation history to list of dicts
    history = []
    if chat_data.conversation_history:
        for msg in chat_data.conversation_history:
            history.append({
                "role": msg.role,
                "content": msg.content
            })
    
    # Get response from GPT
    response = chat_with_gpt(chat_data.message, history)
    
    return ChatResponse(
        message=response,
        role="assistant"
    )

@app.post("/generate-image", response_model=ImageGenerationResponse)
async def generate_image_endpoint(image_data: ImageGenerationRequest):
    """
    Generate an image using DALL-E based on a text prompt.
    """
    if not image_data.prompt.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Prompt cannot be empty"
        )
    
    # Validate size
    valid_sizes = ["1024x1024", "1792x1024", "1024x1792"]
    if image_data.size not in valid_sizes:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid size. Must be one of: {', '.join(valid_sizes)}"
        )
    
    # Validate quality
    valid_qualities = ["standard", "hd"]
    if image_data.quality not in valid_qualities:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid quality. Must be one of: {', '.join(valid_qualities)}"
        )
    
    try:
        image_url = generate_image(
            prompt=image_data.prompt,
            size=image_data.size,
            quality=image_data.quality
        )
        
        return ImageGenerationResponse(
            image_url=image_url,
            prompt=image_data.prompt,
            size=image_data.size
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e)
        )

@app.post("/export")
async def export_summaries(
    export_data: ExportRequest,
    db: Session = Depends(get_db)
):
    """
    Export document summaries in various formats (PDF, TXT, or JSON).
    """
    # Get documents to export
    query = db.query(Document)
    
    if export_data.document_ids:
        query = query.filter(Document.id.in_(export_data.document_ids))
    
    documents = query.order_by(Document.uploaded_at.desc()).all()
    
    if not documents:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No documents found to export"
        )
    
    # Validate format
    valid_formats = ["pdf", "txt", "json"]
    if export_data.format.lower() not in valid_formats:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid format. Must be one of: {', '.join(valid_formats)}"
        )
    
    format_type = export_data.format.lower()
    
    try:
        if format_type == "json":
            # Export as JSON
            export_data_list = []
            for doc in documents:
                export_data_list.append({
                    "filename": doc.filename,
                    "uploaded_at": doc.uploaded_at.isoformat() if doc.uploaded_at else None,
                    "summary": doc.summary or "No summary available",
                    "text_preview": doc.original_text[:500] + "..." if doc.original_text and len(doc.original_text) > 500 else (doc.original_text or "")
                })
            
            json_content = json.dumps({
                "user": "guest",
                "exported_at": datetime.utcnow().isoformat(),
                "total_documents": len(documents),
                "documents": export_data_list
            }, indent=2, ensure_ascii=False)
            
            return Response(
                content=json_content,
                media_type="application/json",
                headers={
                    "Content-Disposition": f'attachment; filename="summaries_{datetime.utcnow().strftime("%Y%m%d_%H%M%S")}.json"'
                }
            )
        
        elif format_type == "txt":
            # Export as TXT
            txt_content = f"Document Summaries Export\n"
            txt_content += f"User: Guest\n"
            txt_content += f"Exported: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}\n"
            txt_content += f"Total Documents: {len(documents)}\n"
            txt_content += "=" * 80 + "\n\n"
            
            for i, doc in enumerate(documents, 1):
                txt_content += f"Document {i}: {doc.filename}\n"
                txt_content += f"Uploaded: {doc.uploaded_at.strftime('%Y-%m-%d %H:%M:%S') if doc.uploaded_at else 'N/A'}\n"
                txt_content += "-" * 80 + "\n"
                txt_content += f"Summary:\n{doc.summary or 'No summary available'}\n"
                txt_content += "\n" + "=" * 80 + "\n\n"
            
            return Response(
                content=txt_content,
                media_type="text/plain",
                headers={
                    "Content-Disposition": f'attachment; filename="summaries_{datetime.utcnow().strftime("%Y%m%d_%H%M%S")}.txt"'
                }
            )
        
        else:
            # Fallback for PDF
            return Response(
                content="PDF export not available in this version",
                media_type="text/plain"
            )

    except Exception as e:
        print(f"Export error: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error exporting documents: {str(e)}"
        )
